#!/usr/bin/env python
import docx
import sys
import json
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import tools

TABLE_RED = "#C00000"
ASSERTION = {1: 'Kém', 3: 'Cần lưu ý', 5: 'Tốt'}

def assert_data(data):
    asserted = {}
    for i in data:
        if i == 'inlet': i = 'temp'
        if i == 'exhaust': continue
        if i == 'raid_stat': continue
        if i == 'mem_util': i = 'mem_free'
        asserted[i] = ['','']
    if data['fault'] == 'No faults found':
        asserted['fault'][0] = 5
        asserted['fault'][1] = 'Lỗi: Không\nĐánh giá: ' + ASSERTION[5]
    else:
        asserted['fault'][0] = 1
        asserted['fault'][1] = 'Lỗi ảnh hưởng tới hoạt động của hệ thống\nĐánh giá: ' + ASSERTION[1]

    temp = data['inlet'].split()[0]
    temp = int(temp)
    if temp >= 21 and temp <= 23:
        asserted['temp'][0] = 5
        asserted['temp'][1] = 'Nhiệt độ bên trong: ' + str(temp) + '\nĐánh giá: ' + ASSERTION[5]
    elif temp > 23 and temp <= 26:
        asserted['temp'][0] = 3
        asserted['temp'][1] = 'Nhiệt độ bên trong: ' + str(temp) + '\nĐánh giá: ' + ASSERTION[3]
    elif temp > 26:
        asserted['temp'][0] = 1
        asserted['temp'][1] = 'Nhiệt độ bên trong: ' + str(temp) + '\nĐánh giá: ' + ASSERTION[1]

    asserted['firmware'][0] = data['firmware'] 
    asserted['firmware'][1] = 'Phiên bản Ilom hiện tại: ' + data['firmware'] + '\nPhiên bản Ilom mới nhất: '

    asserted['image'][0] = data['image']
    asserted['image'][1] = 'Phiên bản OS hiện tại: ' + data['image'] + '\nPhiên bản OS mới nhất:'

    if data['vol_avail'] > 30 and data['raid_stat'] == True:
        asserted['vol_avail'][0] = 5
        asserted['vol_avail'][1] = 'Phân vùng OS được cấu hình RAID \nDung lượng khả dụng: ' + str(data['vol_avail']) + '%' + '\nĐánh giá: ' + ASSERTION[5]
    elif (data['vol_avail'] > 15 and data['vol_avail'] <= 30) and data['raid_stat'] == True:
        asserted['vol_avail'][0] = 3
        asserted['vol_avail'][1] = 'Phân vùng OS được cấu hình RAID \nDung lượng khả dụng: ' + str(data['vol_avail']) + '%' + '\nĐánh giá: ' + ASSERTION[3]
    elif data['vol_avail'] <= 15 and data['raid_stat'] == False:
        asserted['vol_avail'][0] = 1
        asserted['vol_avail'][1] = 'Phân vùng OS không được cấu hình RAID \nDung lượng khả dụng: ' + str(data['vol_avail']) + '%' + '\nĐánh giá: ' + ASSERTION[1]
    elif data['vol_avail'] <= 15 and data['raid_stat'] == True:
        asserted['vol_avail'][0] = 1
        asserted['vol_avail'][1] = 'Phân vùng OS được cấu hình RAID \nDung lượng khả dụng: ' + str(data['vol_avail']) + '%' + '\nĐánh giá: ' + ASSERTION[1]

    if data['bonding'] == 'none':
        asserted['bonding'][0] = 1
        asserted['bonding'][1] = 'Network không được cấu hình bonding'
    elif data['bonding'] == 'aggr':
        asserted['bonding'][0] = 5
        asserted['bonding'][1] = 'Network được cấu hình bonding Aggregration'
    elif data['bonding'] == 'ipmp':
        asserted['bonding'][0] = 5
        asserted['bonding'][1] = 'Network được cấu hình bonding IPMP'

    if data['cpu_util'] <= 30:
        asserted['cpu_util'][0] = 5
    elif data['cpu_util'] > 30 and data['cpu_util'] <= 70:
        asserted['cpu_util'][0] = 3
    else:
        asserted['cpu_util'][0] = 1
    asserted['cpu_util'][1] = 'CPU Ultilization khoảng ' + str(data['cpu_util']) + '%'

    if data['load']['load_avg_per'] <= 2:
        asserted['load'][0] = 5
    elif data['load']['load_avg_per'] > 2 and data['load_avg']['load_avg_per'] <= 5:
        asserted['load'][0] = 3
    else:
        asserted['load'][0] = 1
    asserted['load'][1] = 'CPU Load Average: ' + str(data['load']['load_avg']) + '\nNumber of Cores: ' + str(data['load']['vcpu']) + '\nCPU Load Average per core = CPU Load Average / Number of Cores = ' + str(data['load']['load_avg_per']) 
    
    mem_free = 100 - data['mem_util'] 
    if mem_free >= 20:
        asserted['mem_free'][0] = 5
    elif mem_free > 10 and mem_free < 20:
        asserted['mem_free'][0] = 3
    else:
        asserted['mem_free'][0] = 1
    asserted['mem_free'][1] = 'Physical memory free: ' + str(mem_free) + '%'

    if data['swap_util'] <= 2:
        asserted['swap_util'][0] = 5
    elif data['swap_util'] > 2 and data['swap_util'] <= 5:
        asserted['swap_util'][0] = 3
    else:
        asserted['swap_util'][0] = 1
    asserted['swap_util'][1] = 'SWAP Ultilization: ' + str(data['swap_util']) + '%'
    
    # print(json.dumps(asserted, indent=2))
    return asserted

def get_score(asserted):
    checklist = [
            ['STT', 'Hạng mục kiểm tra', 'Score'],
            [1, 'Kiểm tra trạng thái phần cứng', ['','']],
            [2, 'Kiểm tra nhiệt độ', ['','']],
            [3, 'Kiểm tra phiên bản ILOM', ['','']],
            [4, 'Kiểm tra phiên bản Image', ['','']],
            [5, 'Kiểm tra cấu hình RAID và dung lượng phân vùng OS', ['','']],
            [6, 'Kiểm tra cấu hình Bonding Network', ['','']],
            [7, 'Kiểm tra CPU Utilization', ['','']],
            [8, 'Kiểm tra CPU Load Average', ['','']],
            [9, 'Kiểm tra Memory', ['','']],
            [10, 'Kiểm tra Swap', ['','']],
            ]

    keys = list(asserted)

    for i in range(1, len(checklist)):
        asserted_score = asserted[keys[i-1]][0]
        comment = asserted[keys[i-1]][1]
        try:
            score =  ASSERTION[asserted_score]
        except:
            score = asserted_score
        checklist[i][2][0] = score 
        checklist[i][2][1] = comment

    # print(checklist)
    return checklist

# #!@#!@#@!
def drw_table(doc, checklist, row, col, info=False):
    if checklist == []:
        return -1
    tab = doc.add_table(row, col)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # ADD TITLE CELLS AND COLOR THEM 
    cols = tab.rows[0].cells
    for r in range(len(checklist[0])):
        cell = cols[r]
        cell.text = checklist[0][r]
        cell.paragraphs[0].style = 'Table Heading'
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), TABLE_RED))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # ADD CONTENT TO TABLE
    if not info:
        for i in range(1, len(checklist)):
            rows = tab.rows[i]
            cells = rows.cells
            for j in range(0, len(checklist[i])):
                cells[j].text = checklist[i][j]
                cells[j].paragraphs[0].style = 'Table Contents'
    else:
        for i in range(1, len(checklist)):
            rows = tab.rows[i]
            cells = rows.cells
            for j in range(0, len(checklist[i])):
                if j == (len(checklist[i])-1):
                    cells[j].text = checklist[i][j][0]
                    cells[j].paragraphs[0].style = 'Table Contents'
                    continue
                cells[j].text = str(checklist[i][j])
                cells[j].paragraphs[0].style = 'Table Contents'
    return tab


def drw_info(doc, node, checklist, images=[]):
    print(node)
    for i in range(1, len(checklist)):
        doc.add_paragraph(checklist[i][1], style='baocao4')
        if isinstance(images[i-1], list):
            for image in images[i-1]:
                path = './output/' + node + '/' + image
                doc.add_picture(path, width=Inches(6.73))
        else:
            doc.add_picture('./output/' + node + '/' + images[i-1], width=Inches(6.73))
        doc.add_paragraph(checklist[i][2][1])
    doc.add_page_break()

def define_doc():
    try:
        doc = docx.Document("./sample/dcx8m2.docx")
    except Exception as err:
        print('Error:', err)
        sys.exit()
    return doc

def drw_doc(doc, nodes):
    data = tools.read_json(nodes + '.json')
    asserted_list = []
    for node in data:
        # print('TYPE OF NOTE IN DATA ', end='')
        # print(type(node))
        # print(json.dumps(node, indent=2))
        images = tools.read_json('./output/' + node + '/images.json')
        print(json.dumps(images, indent=2))
        file_dump = {}
        asserted = assert_data(data[node])
        
        file_dump[node] = asserted

        keys = list(asserted)
        checklist = get_score(asserted)
        doc.add_paragraph("Máy chủ " + node, style='baocao2')
        doc.add_paragraph("Thông tin tổng quát", style='baocao3')
        overview = [
                ['Hostname', 'Product Name', 'Serial Number', 'IP Address'],
                [node,'','',''],
                ]
        drw_table(doc, overview, 2, 4)
        doc.add_paragraph("Đánh giá", style='baocao3')
        drw_table(doc, checklist, 11, 3, True)

        doc.add_paragraph("Thông tin chi tiết", style='baocao3')
        drw_info(doc, node, checklist, images)
        asserted_file = node + '_asserted'
        asserted_list += [asserted_file]
        
        tools.save_json('./output/' + node + '/' + asserted_file, file_dump)
    tools.join_json(asserted_list, nodes + '_asserted')
    print()
    return doc

def print_style(doc):
    styles = doc.styles
    p_styles = [
            s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
            ]
    for style in p_styles:
        print(style.name)

def run(output):
    doc = define_doc() 
    doc = drw_doc(doc, output)

    print_style(doc)
    doc.save(output + '.docx')

##### MAIN #####
def main():
    rc = 1
    run()

if __name__ == "__main__":
    main()

