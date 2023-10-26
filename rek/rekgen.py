import os, sys
import shutil, glob, re
import io 
import zipfile, tarfile
import argparse
import json
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

TABLE_RED = "#C00000"
ASSERTION = {1: 'Kém', 3: 'Cần lưu ý', 5: 'Tốt'}

##### DEFAULT PATHS #####
## INTERGRATED LIGHT OUT MANAGEMENT
FAULT='/fma/@usr@local@bin@fmadm_faulty.out'
TEMP='/ilom/@usr@local@bin@collect_properties.out'
FIRMWARE='/ilom/@usr@local@bin@collect_properties.out'
##
## ORACLE LINUX
IMAGE_LINUX=''
PARTITION_LINUX='/disks/df-kl.out'
RAID='' 
NETWORK='/sysconfig/ifconfig-a.out'
CPU_ULTILIZATION=''
CPU_LOAD_LINUX=''
MEM_LINUX=''
SWAP_SOL=''
EXTRACT_LOCATION='./temp/'
#HugePages = HugePages_Total * 2 /1024 = ~ 67.8% physical Memory
##
## ORACLE SOLARIS
IMAGE_SOL='/etc/release'
PARTITION_SOL='/disks/df-kl.out'
RAID_SOL='/disks/zfs/zpool_status_-v.out' 
NETWORK_SOL='/netinfo/ipadm.out'
CPU_ULTILIZATION_SOL='/sysconfig/vmstat_3_3.out'
CPU_LOAD_SOL='/sysconfig/prstat-L.out'
VCPU_SOL='/ldom/ldm_list.out'
MEM_SOL='/disks/zfs/mdb/mdb-memstat.out'
SWAP_SOL='/disks/swap-s.out'
##
##### END_PATHS #####

##### JSON #####
#get a dictionary as input and dump it to a json type file
def save_json(file, content):
    if not content:
        print('No content from input to save!')
        return -1

    try:
        with open(file, 'w') as f:
            json.dump(content, f, indent=2)
    except OSError as err:
        print('OS error: ', err)
        raise RuntimeError('Cannot save JSON') from err
        return -1

def read_json(file):
    try:
        with open(file, 'r+') as f:
            content = json.load(f)
        return content
    except OSError as err:
        print('OS error: ', err)
        raise RuntimeError('Cannot read JSON') from err
        return -1

def join_json(output):
    try:
        with open('./output/data.json', 'a+') as file:
            x = {}
            for i in output:
                path = './output/' + i + '.json'
                buffer = read_json(path)
                key = list(buffer)[0]
                x[key] = buffer[key]
            json.dump(x, file, indent=4)
    except OSError as err:
        print('OS error: ', err)
        return -1

##### END JSON #####

def save_file(file, content):
    try:
        with open(file, 'w') as f:
            f.write(content)
    except OSError as err:
        raise RuntimeError('Cannot save file: ') from err
        return -1

def rm_ext(file, ext):
    return file.split('/')[2][:-len(ext)-1]

##### IMPLEMETATION #####
def clean_files(folder='./temp/'):
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        print(file_path)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))
            return -1

def clean_up(path='./temp/'):
    print('Remove unzip files? (y/n) ', end='')
    choice = input()
    if choice in ['', 'yes', 'y', 'Y', 'yeah', 'YES']:
        clean_files(path)
        return
    return -1 

def clean_up_force(path='./temp/'):
    print('FORCE CLEAN UP DUE TO ERROR!')
    clean_files(path)
    return -1

def check_valid(path):
    return os.path.isdir(path)

def extract_file(serial, compress):
    compress = compress.lower()
    # regex = '*[_.]' + serial + '[_.]*.' + compress
    regex = '*' + serial + '*.' + compress
    file = get_file(regex, root='./sample/') 
    if file == -1: return -1

    print('Extracting: ', file)
    if compress == 'zip':
        unzip(file)
        return rm_ext(file, compress)
    elif compress == 'tar.gz':
        untar(file)
        return rm_ext(file, compress)
    else: return -1

# Find return the file with serial number 
def get_file(regex, root=''):
    path = root + regex
    files = glob.glob(path, recursive=True)
    if len(files) == 0:
        print('No file found matched!')
        return -1
    elif len(files) == 1:
        return files[0]
    else:
        for i in range(len(files)):
            print('[', i, '] ', files[i], sep='')
        while True:
            try:
                c = int(input('Which file?\n [0] ') or '0')
            except:
                continue
            if c < 0 and c > len(files):
                continue
            else: 
                break
        return files[c]

def grep(path, regex, single_line=True):
    result = ''
    flag = re.MULTILINE
    pattern = re.compile(regex, flag)
    file = cat(path, False)
    content = file.readlines()

    if single_line:
        for line in range(len(content)):
            if re.search(pattern, content[line]):
                result += content[line].lstrip()
                print(line + 1, ': ', content[line], sep='', end='')
                break 
    else:
        for line in range(len(content)):
            if re.search(pattern, content[line]):
                result += content[line].lstrip()
                print(line + 1, ': ', content[line], sep='', end='')
    print()
    return result

def get_ilom(path):
    print('##### ILOM #####')
    fault = cat(path + FAULT).strip()

    inlet_temp = grep(path + TEMP, 'inlet_temp').strip().split()
    inlet_temp = ' '.join(inlet_temp[2:5])

    exhaust_temp = grep(path + TEMP, 'exhaust_temp').strip().split()
    exhaust_temp = ' '.join(exhaust_temp[2:5])

    firmware = grep(path + FIRMWARE, 'Version').strip().split()
    firmware = ' '.join(firmware[1:])

    print('##### END OF ILOM #####')
    return {'fault': fault, 'inlet': inlet_temp, 'exhaust': exhaust_temp, 'firmware': firmware} 

def get_os(path, os='SOL'):
    x = {}
    if os == 'SOL':
        image = grep(path + IMAGE_SOL, 'Solaris').strip().split()
        image = image[2]
        x['image'] = image

        vol = grep(path + PARTITION_SOL, "\\B\/\\B").strip().split()
        vol = vol[-2]
        vol_avail = 100 - int(vol[:-1])
        x['vol_avail'] = vol_avail

        raid = grep(path + RAID_SOL, "mirror").strip().split()
        if 'ONLINE' in raid:
            raid_stat = True 
        else:
            raid_stat = False
        x['raid_stat'] = raid_stat

        net_ipmp = grep(path + NETWORK_SOL, 'ipmp')
        net_aggr = grep(path + NETWORK_SOL, 'aggr')
        if not net_ipmp and not net_aggr:
            bonding = 'none'
        elif net_ipmp and not net_aggr:
            bonding = 'ipmp'
        elif net_aggr and not net_ipmp:
            bonding = 'aggr'
        else:
            bonding = 'both'
        x['bonding'] = bonding

        cpu_idle = cat(path + CPU_ULTILIZATION_SOL).strip().split('\n')
        cpu_idle = cpu_idle[2]
        cpu_idle = cpu_idle.split()[21]
        cpu_util = 100 - int(cpu_idle)
        x['cpu_util'] = cpu_util

        x['load'] = {}
        load = grep(path + CPU_LOAD_SOL, 'load average').strip().split(', ')
        load_avg = ' '.join(load).split()[-3:]
        load_avg = float((max(load_avg)))
        vcpu = grep(path + VCPU_SOL, 'primary').strip().split()[4]
        load_avg_per = load_avg / float(vcpu)
        load_avg_per = float(f'{load_avg_per:.3f}')
        x['load']['load_avg'] = load_avg
        x['load']['vcpu'] = vcpu
        x['load']['load_avg_per'] = load_avg_per

        mem = grep(path + MEM_SOL, 'freelist', False).strip().split()
        mem_free = mem[-1]
        mem_util = 100 - int(mem_free[:-1])
        x['mem_util'] = mem_util

        swap = cat(path + SWAP_SOL).strip().split()
        swap = [swap[8], swap[10]]
        swap[0] = int(swap[0][:-2])
        swap[1] = int(swap[1][:-2])
        swap_util = swap[0] / (swap[0] + swap[1])
        swap_util = int(swap_util * 100)
        x['swap_util'] = swap_util

        print(x)
        print()
    return x

def get_content(path):
    root = './temp/'
    orj_path = path[0]
    for i in range(0, len(path)):
        path[i] = root + str(path[i])

    # @@
    ilom = get_ilom(path[0])
    os_info = get_os(path[1], 'SOL')
    name = orj_path.split('_')[0]

    content = {}
    content[name] = {
            'fault': ilom['fault'],
            'inlet': ilom['inlet'],
            'exhaust': ilom['exhaust'],
            'firmware': ilom['firmware'],
            'image': os_info['image'],
            'vol_avail': os_info['vol_avail'],
            'raid_stat': os_info['raid_stat'],
            'bonding': os_info['bonding'],
            'cpu_util': os_info['cpu_util'],
            'load': os_info['load'],
            'mem_util': os_info['mem_util'],
            'swap_util': os_info['swap_util'],
    }
    print('##### NODE INFORMATION #####')
    print(json.dumps(content, indent = 2))
    print('##### END OF NODE INFORMATION #####\n')
    return content

def cat(path, stdout=True):
    result = ''
    try:
        with open(path, 'r') as file:
            content = file.readlines()
            if stdout:
                count = 0
                for l in content:
                    result += l.lstrip()
                    print(++count + 1, ': ', l, sep='', end='')
                return result
            else: 
                for l in content:
                    result += l.lstrip()
                return io.StringIO(result)
    except Exception as err:
        raise RuntimeError('Cannot open file to read') from err
        return -1

def unzip(file):
    if not zipfile.is_zipfile(file):
        print('Error: Not a zip file')
        return -1
    try:
        with zipfile.ZipFile(file, 'r') as z_object:
            z_object.extractall(path='./temp/')
            print('> UNZIP:', file)
    except Exception as err:
        print('Error:' , err)
        return -1

def untar(file):
    # sucks
    if not tarfile.is_tarfile(file):
        print('Error: Not a tar file')
        return -1
    print(file)
    try:
        with tarfile.open(file, 'r') as t_object:
            try: 
                t_object.extractall(path='./temp/')
                print('> UNTAR:', file)
            except:
                buffer = rm_ext(file, 'tar.gz')
                if clean_up('./temp/' + buffer) == -1:
                    return -1
                t_object.extractall(path='./temp/')
    except Exception as err:
        print(err)
        return -1

def compile():
    try:
        with open('./input', 'r') as file:
            number = file.readlines()
    except Exception as err:
        print(err)
        return -1

    output_files = []
    for i in range(len(number)):
        serial = number[i].strip()
        print(serial)
        # if len(serial) != 1: 
        #     print('Error: Only one name each line!')
        #     return -1

        path = ['','']
        print('##### EXTRACT FILES #####')
        path[0] = extract_file(serial, 'zip')
        path[1] = extract_file(serial, 'tar.gz')
        print('##### END EXTRACTION #####\n')

        if path == [-1, -1]: 
            print('Error: PATH not exist!')
            return -1

        print('PATH: ', path)

        data = input('Output file: ').strip()
        output_files += [data]

        data = './output/' + data + '.json'
        content = get_content(path)
        if save_json(data, content) == -1:
            return -1 
    return output_files

##### END_IMPLEMENTATION #####


def assert_data(data):
    asserted = {}
    for i in data:
        if i == 'inlet': i = 'temp'
        if i == 'exhaust': continue
        if i == 'raid_stat': continue
        if i == 'mem_util': i = 'mem_free'
        asserted[i] = ['','']
    if data['fault'] == 'No faults found':
        asserted['fault'][0] = 5
        asserted['fault'][1] = 'Lỗi: Không\nĐánh giá: ' + ASSERTION[5]
    else:
        asserted['fault'][0] = 1
        asserted['fault'][1] = 'Lỗi ảnh hưởng tới hoạt động của hệ thống\nĐánh giá: ' + ASSERTION[1]

    temp = data['inlet'].split()[0]
    temp = int(temp)
    if temp >= 21 and temp <= 23:
        asserted['temp'][0] = 5
        asserted['temp'][1] = 'Nhiệt độ bên trong: ' + str(temp) + '\nĐánh giá: ' + ASSERTION[5]
    elif temp > 23 and temp <= 26:
        asserted['temp'][0] = 3
        asserted['temp'][1] = 'Nhiệt độ bên trong: ' + str(temp) + '\nĐánh giá: ' + ASSERTION[3]
    elif temp > 26:
        asserted['temp'][0] = 1
        asserted['temp'][1] = 'Nhiệt độ bên trong: ' + str(temp) + '\nĐánh giá: ' + ASSERTION[1]

    asserted['firmware'][0] = data['firmware'] 
    asserted['firmware'][1] = 'Phiên bản Ilom hiện tại: ' + data['firmware'] + '\nPhiên bản Ilom mới nhất: '

    asserted['image'][0] = data['image']
    asserted['image'][1] = 'Phiên bản OS hiện tại: ' + data['image'] + '\nPhiên bản OS mới nhất:'

    if data['vol_avail'] > 30 and data['raid_stat'] == True:
        asserted['vol_avail'][0] = 5
        asserted['vol_avail'][1] = 'Phân vùng OS được cấu hình RAID \nDung lượng khả dụng: ' + str(data['vol_avail']) + '%' + '\nĐánh giá: ' + ASSERTION[5]
    elif (data['vol_avail'] > 15 and data['vol_avail'] <= 30) and data['raid_stat'] == True:
        asserted['vol_avail'][0] = 3
        asserted['vol_avail'][1] = 'Phân vùng OS được cấu hình RAID \nDung lượng khả dụng: ' + str(data['vol_avail']) + '%' + '\nĐánh giá: ' + ASSERTION[3]
    elif data['vol_avail'] <= 15 and data['raid_stat'] == False:
        asserted['vol_avail'][0] = 1
        asserted['vol_avail'][1] = 'Phân vùng OS không được cấu hình RAID \nDung lượng khả dụng: ' + str(data['vol_avail']) + '%' + '\nĐánh giá: ' + ASSERTION[1]
    elif data['vol_avail'] <= 15 and data['raid_stat'] == True:
        asserted['vol_avail'][0] = 1
        asserted['vol_avail'][1] = 'Phân vùng OS được cấu hình RAID \nDung lượng khả dụng: ' + str(data['vol_avail']) + '%' + '\nĐánh giá: ' + ASSERTION[1]

    if data['bonding'] == 'none':
        asserted['bonding'][0] = 1
        asserted['bonding'][1] = 'Network không được cấu hình bonding'
    elif data['bonding'] == 'aggr':
        asserted['bonding'][0] = 5
        asserted['bonding'][1] = 'Network được cấu hình bonding Aggregration'
    elif data['bonding'] == 'ipmp':
        asserted['bonding'][0] = 5
        asserted['bonding'][1] = 'Network được cấu hình bonding IPMP'

    if data['cpu_util'] <= 30:
        asserted['cpu_util'][0] = 5
    elif data['cpu_util'] > 30 and data['cpu_util'] <= 70:
        asserted['cpu_util'][0] = 3
    else:
        asserted['cpu_util'][0] = 1
    asserted['cpu_util'][1] = 'CPU Ultilization khoảng ' + str(data['cpu_util']) + '%'

    if data['load']['load_avg_per'] <= 2:
        asserted['load'][0] = 5
    elif data['load']['load_avg_per'] > 2 and data['load_avg']['load_avg_per'] <= 5:
        asserted['load'][0] = 3
    else:
        asserted['load'][0] = 1
    asserted['load'][1] = 'CPU Load Average: ' + str(data['load']['load_avg']) + '\nNumber of Cores: ' + data['load']['vcpu'] + '\nCPU Load Average per core = CPU Load Average / Number of Cores = ' + str(data['load']['load_avg_per']) 
    
    mem_free = 100 - data['mem_util'] 
    if mem_free >= 20:
        asserted['mem_free'][0] = 5
    elif mem_free > 10 and mem_free < 20:
        asserted['mem_free'][0] = 3
    else:
        asserted['mem_free'][0] = 1
    asserted['mem_free'][1] = 'Physical memory free: ' + str(mem_free) + '%'

    if data['swap_util'] <= 2:
        asserted['swap_util'][0] = 5
    elif data['swap_util'] > 2 and data['swap_util'] <= 5:
        asserted['swap_util'][0] = 3
    else:
        asserted['swap_util'][0] = 1
    asserted['swap_util'][1] = 'SWAP Ultilization: ' + str(data['swap_util']) + '%'
    
    print(json.dumps(asserted, indent=2))
    return asserted

def get_score(asserted):
    checklist = [
            ['STT', 'Hạng mục kiểm tra', 'Score'],
            [1, 'Kiểm tra trạng thái phần cứng', ['','']],
            [2, 'Kiểm tra nhiệt độ', ['','']],
            [3, 'Kiểm tra phiên bản ILOM', ['','']],
            [4, 'Kiểm tra phiên bản Image', ['','']],
            [5, 'Kiểm tra cấu hình RAID và dung lượng phân vùng OS', ['','']],
            [6, 'Kiểm tra cấu hình Bonding Network', ['','']],
            [7, 'Kiểm tra CPU Utilization', ['','']],
            [8, 'Kiểm tra CPU Load Average', ['','']],
            [9, 'Kiểm tra Memory', ['','']],
            [10, 'Kiểm tra Swap', ['','']],
            ]

    keys = list(asserted)

    for i in range(1, len(checklist)):
        asserted_score = asserted[keys[i-1]][0]
        comment = asserted[keys[i-1]][1]
        try:
            score =  ASSERTION[asserted_score]
        except:
            score = asserted_score
        checklist[i][2][0] = score 
        checklist[i][2][1] = comment

    print(checklist)
    return checklist

# #!@#!@#@!
def drw_table(doc, checklist, row, col, info=False):
    if checklist == []:
        return -1
    tab = doc.add_table(row, col)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # ADD TITLE CELLS AND COLOR THEM 
    cols = tab.rows[0].cells
    for r in range(len(checklist[0])):
        cell = cols[r]
        cell.text = checklist[0][r]
        cell.paragraphs[0].style = 'Table Heading'
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), TABLE_RED))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # ADD CONTENT TO TABLE
    if not info:
        for i in range(1, len(checklist)):
            rows = tab.rows[i]
            cells = rows.cells
            for j in range(0, len(checklist[i])):
                cells[j].text = checklist[i][j]
                cells[j].paragraphs[0].style = 'Table Contents'
    else:
        for i in range(1, len(checklist)):
            rows = tab.rows[i]
            cells = rows.cells
            for j in range(0, len(checklist[i])):
                if j == (len(checklist[i])-1):
                    cells[j].text = checklist[i][j][0]
                    cells[j].paragraphs[0].style = 'Table Contents'
                    continue
                cells[j].text = str(checklist[i][j])
                cells[j].paragraphs[0].style = 'Table Contents'
    return tab

def drw_info(doc, checklist):
    for i in range(1, len(checklist)):
        doc.add_paragraph(checklist[i][1], style='baocao4')
        doc.add_paragraph(checklist[i][2][1])
    doc.add_page_break()

def define_doc():
    try:
        doc = docx.Document("./sample/dcx8m2.docx")
    except Exception as err:
        print('Error:', err)
        sys.exit()
    return doc

def drw_doc(doc):
    data = read_json('./output/data.json')
    print(json.dumps(data, indent=2))
    for node in data:
        file_dump = {}
        asserted = assert_data(data[node])
        file_dump[node] = asserted

        keys = list(asserted)
        checklist = get_score(asserted)
        doc.add_paragraph("Máy chủ " + node, style='baocao2')

        doc.add_paragraph("Thông tin tổng quát", style='baocao3')
        overview = [
                ['Hostname', 'Product Name', 'Serial Number', 'IP Address'],
                [node,'','',''],
                ]
        drw_table(doc, overview, 2, 4)
        doc.add_paragraph("Đánh giá", style='baocao3')
        drw_table(doc, checklist, 11, 3, True)

        doc.add_paragraph("Thông tin chi tiết", style='baocao3')
        drw_info(doc, checklist)
        save_json( './output/' + node + '_asserted.json' ,file_dump)

    print()
    return doc

def print_style(doc):
    styles = doc.styles
    p_styles = [
            s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
            ]
    for style in p_styles:
        print(style.name)

def doc_run():
    doc = define_doc() 
    doc = drw_doc(doc)

    print_style(doc)
    doc.save("./output/dcx8m2_out.docx")

# FLOW OF PROGRAM
def run():
    output_files = compile()
    if output_files == -1:
        print('Error: No files to join!')
        return -1

    choice = input('Join all input?[y/n] ')
    if choice in ['', 'yes', 'y', 'Y', 'yeah', 'YES']:
        join_json(output_files)

    choice = input('GENERATE DOCUMENT?[y/n] ')
    if choice in ['', 'yes', 'y', 'Y', 'yeah', 'YES']:
        doc_run()

##### MAIN #####
def main():
    if run() == -1: 
        clean_up_force()
        return -1
    clean_up()
    sys.exit()

if __name__ == "__main__":
    main()
##### END_MAIN #####
