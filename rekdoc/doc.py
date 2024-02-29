#!/usr/bin/env python
###
import sys
import os
# import core.logger
import json
import docx
from pathlib import Path
from rekdoc import core

#
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document

#
from rekdoc import tools

TABLE_RED = "#C00000"
ASSERTION = {1: "Kém", 3: "Cần lưu ý", 5: "Tốt"}


def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def assert_fault(data: dict) -> list:
    score = 0
    # if data["fault"] == "":
    #     score = ""
    #     comment = [""]
    if data["fault"] == "No faults found" or data["fault"] == "":
        score = 5
        comment = ["Lỗi: Không", "Đánh giá: " + ASSERTION[score]]
    elif data["fault"] == "warning":
        score = 3
        comment = [
            "Lỗi không ảnh hưởng tới hiệu năng của hệ thống",
            "Đánh giá: " + ASSERTION[score],
        ]
    else:
        score = 1
        comment = [
            "Lỗi ảnh hưởng tới hoạt động của hệ thống",
            "Đánh giá: " + ASSERTION[score],
        ]

    fault = [score, comment]
    core.logger.debug(json.dumps(fault, ensure_ascii=False))
    return fault


def assert_temp(data: dict) -> list:
    if data["inlet"] == "":
        score = ""
        comment = [""]
        temp = [score, comment]
        return temp
    inlet_temp = data["inlet"].split()[0]
    inlet_temp = int(inlet_temp)
    score = 0
    comment = [""]
    if inlet_temp <= 23:
        score = 5
        comment = [
            "Nhiệt độ bên trong: " + str(inlet_temp),
            "Đánh giá: " + ASSERTION[score],
        ]
    elif inlet_temp > 23 and inlet_temp <= 26:
        score = 3
        comment = [
            "Nhiệt độ bên trong: " + str(inlet_temp),
            "Đánh giá: " + ASSERTION[score],
        ]
    elif inlet_temp > 26:
        score = 1
        comment = [
            "Nhiệt độ bên trong: " + str(inlet_temp),
            "Đánh giá: " + ASSERTION[score],
        ]

    temp = [score, comment]
    core.logger.debug(json.dumps(temp, ensure_ascii=False))
    return temp


def assert_firmware(data: dict) -> list:
    latest = ""
    if data["firmware"] == "":
        firmware = ["", [""]]
        core.logger.debug(json.dumps(firmware, ensure_ascii=False))
        return firmware
    while True:
        try:
            sys.stdout.write("\033[?25h")
            latest = (
                input(
                    "Enter latest ILOM version\n[" + data["firmware"] + "] "
                )
                or data["firmware"]
            )
        except KeyboardInterrupt:
            print()
            sys.exit()
        except ValueError:
            continue
        break

    if latest == data["firmware"]:
        score = 5
    else:
        while True:
            try:
                print("Đánh giá")
                print("[0] Tốt")
                print("[1] Cần lưu ý")
                print("[2] Kém")
                choice = int(input("Chọn đánh giá\n [0] ") or "0")
                if choice == 0:
                    score = 5
                elif choice == 1:
                    score = 3
                else:
                    score = 1
            except KeyboardInterrupt:
                print()
                sys.exit()
            except ValueError:
                continue
            break

    comment = [
        "Phiên bản ILOM hiện tại: " + data["firmware"],
        "Phiên bản ILOM mới nhất: " + latest,
        "Đánh giá: " + ASSERTION[score]
    ]
    firmware = [score, comment]
    core.logger.debug(json.dumps(firmware, ensure_ascii=False))
    return firmware


def assert_image(data: dict) -> list:
    score = 0
    if data["image"] == "":
        image = [score, [""]]
        return image

    while True:
        try:
            sys.stdout.write("\033[?25h")
            latest = (
                input("Enter latest OS version\n[" + data["image"] + "] ")
                or data["image"]
            )
        except KeyboardInterrupt:
            print()
            sys.exit()
        except ValueError:
            continue
        break

    if latest == data["image"]:
        score = 5
    else:
        while True:
            try:
                print("Đánh giá")
                print("[0] Tốt")
                print("[1] Cần lưu ý")
                print("[2] Kém")
                choice = int(input("Chọn đánh giá\n [0] ") or "0")
                if choice == 0:
                    score = 5
                elif choice == 1:
                    score = 3
                else:
                    score = 1
            except KeyboardInterrupt:
                print()
                sys.exit()
            except ValueError:
                continue
            break

    comment = [
        "Phiên bản OS hiện tại: " + data["image"],
        "Phiên bản OS mới nhất: " + latest,
        "Đánh giá: " + ASSERTION[score]
    ]
    image = [score, comment]
    core.logger.debug(json.dumps(image, ensure_ascii=False))
    return image


def assert_vol(data: dict) -> list:
    score = 0
    comment = [""]
    if data["vol_avail"] == "":
        score = ""
        comment = [""]
        vol = [score, comment]
        return vol
    elif system_info["type"] == "vm":
        if data["vol_avail"] > 30:
            score = 5
        elif (data["vol_avail"] > 15 and data["vol_avail"] <= 30):
            score = 3
        elif data["vol_avail"] <= 15:
            score = 1
    else:
        if data["vol_avail"] > 30 and data["raid_stat"] is True:
            score = 5
            comment = [
                "Phân vùng OS được cấu hình RAID",
            ]
        elif (data["vol_avail"] > 15 and data["vol_avail"] <= 30) and data[
            "raid_stat"
        ] is True:
            score = 3
            comment = [
                "Phân vùng OS được cấu hình RAID",
            ]
        elif data["vol_avail"] <= 15 and data["raid_stat"] is False:
            score = 1
            comment = [
                "Phân vùng OS không được cấu hình RAID",
            ]
        elif data["vol_avail"] <= 15 and data["raid_stat"] is True:
            score = 1
            comment = [
                "Phân vùng OS được cấu hình RAID",
            ]
        elif data["vol_avail"] > 30 and data["raid_stat"] is False:
            score = 3
            comment = [
                "Phân vùng OS không được cấu hình RAID",
            ]

    comment.extend(["Dung lượng khả dụng: " +
                    str(data["vol_avail"]) + "%",
                    "Đánh giá: " + ASSERTION[score]])

    vol = [score, comment]
    core.logger.debug(json.dumps(vol, ensure_ascii=False))

    return vol


def assert_bonding(data: dict) -> list:
    if data["bonding"] == "":
        score = ""
        comment = [""]
        bonding = [score, comment]
        return bonding
    elif data["bonding"] == "none":
        score = 1
        comment = ["Network không được cấu hình bonding"]
    elif data["bonding"] == "aggr":
        score = 5
        comment = ["Network được cấu hình bonding Aggregration"]
    elif data["bonding"] == "ipmp":
        score = 5
        comment = ["Network được cấu hình bonding IPMP"]
    comment.append("Đánh giá: " + ASSERTION[score])

    bonding = [score, comment]
    core.logger.debug(json.dumps(bonding, ensure_ascii=False))

    return bonding


def assert_cpu_util(data: dict) -> list:
    if data["cpu_util"] == "":
        score = ""
        comment = [""]
        cpu_util = [score, comment]
        return cpu_util
    elif data["cpu_util"] <= 30:
        score = 5
    elif data["cpu_util"] > 30 and data["cpu_util"] <= 70:
        score = 3
    else:
        score = 1
    comment = ["CPU Utilization khoảng " + str(data["cpu_util"]) + "%"]
    comment.append("Đánh giá: " + ASSERTION[score])

    cpu_util = [score, comment]
    core.logger.debug(json.dumps(cpu_util, ensure_ascii=False))

    return cpu_util


def assert_load(data: dict) -> list:
    if data["load_avg"] == "":
        score = ""
        comment = [""]
    elif data["load"]["load_avg_per"] <= 2:
        score = 5
    elif 2 < data["load"]["load_avg_per"] <= 5:
        score = 3
    else:
        score = 1
    comment = [
        "CPU Load Average: " + str(data["load"]["load_avg"]),
        "Number of Cores: " + str(data["load"]["vcpu"]),
        "CPU Load Average per core = CPU Load Average / Number of Cores = "
        + str(data["load"]["load_avg_per"]),
        "Đánh giá: " + ASSERTION[score]
    ]

    load = [score, comment]
    core.logger.debug(json.dumps(load, ensure_ascii=False))

    return load


def assert_mem_free(data: dict) -> list:
    # mem_free = 100 - data["mem_util"]
    mem_free = data["mem_free"]["mem_free_percent"]
    if mem_free == "":
        score = ""
        comment = [""]
        mem_free = [score, comment]
        core.logger.debug(json.dumps(mem_free, ensure_ascii=False))
        return mem_free
    elif mem_free >= 20:
        score = 5
    elif mem_free > 10 and mem_free < 20:
        score = 3
    else:
        score = 1
    comment = ["Total Memory: " + str(data["mem_free"]["total_mem"])]
    comment.append("Memory Free in GB: " + str(data["mem_free"]["mem_free"]))
    comment.append("Average physical memory free: " + str(mem_free) + "%")
    comment.append("Đánh giá: " + ASSERTION[score])

    mem_free = [score, comment]
    core.logger.debug(json.dumps(mem_free, ensure_ascii=False))

    return mem_free


def assert_io_busy(data: dict) -> list:
    score = 0
    comment = []
    if not data["io_busy"]:
        score = ""
        comment = [""]
        io_busy = [score, comment]
        return io_busy
    elif data["io_busy"]["busy"] < 50:
        score = 5
        comment = ["IO Busy: " + "< 50%"]
    elif 50 <= data["io_busy"]["busy"] <= 70:
        score = 3
        comment = ["IO Busy: " + ">= 50%" + " " + "và" + "<= 70%"]
    else:
        score = 1
        comment = ["IO Busy: " + "> 70%"]

    # comment = ["Thiết bị IO Busy cao: " + data["io_busy"]["name"]]
    # comment.append("IO Busy: " + str(data["io_busy"]["busy"]))
    comment.append("Đánh giá: " + ASSERTION[score])
    io_busy = [score, comment]
    return io_busy


def assert_swap_util(data: dict) -> list:
    if data["swap_util"] <= 2:
        score = 5
    elif data["swap_util"] > 2 and data["swap_util"] <= 5:
        score = 3
    else:
        score = 1
    comment = ["SWAP Utilization: " + str(data["swap_util"]) + "%"]

    swap_util = [score, comment]
    core.logger.debug(json.dumps(swap_util, ensure_ascii=False))

    return swap_util


def assert_ilom(data: dict) -> dict:
    x = {}
    try:
        fault = assert_fault(data)
        temp = assert_temp(data)
        firmware = assert_firmware(data)
        x = {"fault": fault,
             "temp": temp,
             "firmware": firmware}
    except RuntimeError:
        print("Failed to assert ILOM")
        raise
    core.logger.debug(json.dumps(x, indent=2))
    return x


def assert_system_status(data: dict, server_type: str) -> dict:
    x = {}
    try:
        image = assert_image(data)
        vol = assert_vol(data)
        x = {"image": image,
             "vol": vol}
        if server_type == "baremetal":
            bonding = assert_bonding(data)
            x["bonding"] = bonding
    except RuntimeError:
        print("Failed to assert system status")
        raise
    return x


def assert_system_perform(data: dict, platform: str, system_type: str) -> dict:
    x = {}
    try:
        if system_type == "standalone":
            if platform == "solaris":
                cpu_util = assert_cpu_util(data)
                mem_free = assert_mem_free(data)
                io_busy = assert_io_busy(data)

                x = {"cpu_util": cpu_util,
                     "mem_free": mem_free,
                     "io_busy": io_busy}
            elif platform == "linux":
                pass
        elif system_type == "exa":
            pass
    except RuntimeError:
        print("Failed to assert system performance")
        raise
    core.logger.debug(json.dumps(x))
    return x


def assert_data(data: dict) -> dict:
    asserted = {}
    ilom = {}

    if system_info["type"] == "baremetal":
        ilom = assert_ilom(data)
    else:
        ilom = {"fault": assert_fault(data)}
    system_status = None
    system_perform = None
    if system_info["system_type"] == "standalone":
        system_status = assert_system_status(data,
                                             system_info["type"])
        system_perform = assert_system_perform(data,
                                               system_info["platform"],
                                               system_info["system_type"])
    elif system_info["system_type"] == "exa":
        system_status = assert_system_status(data,
                                             system_info["type"])
        system_perform = assert_system_perform(data,
                                               system_info["platform"],
                                               system_info["system_type"])
    asserted = {"node_name": data["node_name"],
                **ilom,
                **system_status,
                **system_perform}

    # cpu_util = assert_cpu_util(data)
    # asserted["cpu_util"][0] = cpu_util[0]
    # asserted["cpu_util"][1].extend(cpu_util[1])
    #
    # mem_free = assert_mem_free(data)
    # asserted["mem_free"][0] = mem_free[0]
    # asserted["mem_free"][1].extend(mem_free[1])
    # load = assert_load(data)
    # asserted["load"][0] = load[0]
    # asserted["load"][1].extend(load[1])
    #
    #
    # swap_util = assert_swap_util(data)
    # asserted["swap_util"][0] = swap_util[0]
    # asserted["swap_util"][1].extend(swap_util[1])

    for field in asserted:
        core.logger.debug("ASSERTED:" + field + ": " + str(asserted[field][0]))
    return asserted


def get_score(asserted: dict) -> list:
    checklist = []
    if system_info["type"] == "baremetal" \
            and system_info["platform"] == "solaris":
        checklist = [
            # ["STT", "Hạng mục kiểm tra", "Score"],
            [1, "Kiểm tra trạng thái phần cứng", ["", []]],
            [2, "Kiểm tra nhiệt độ", ["", []]],
            [3, "Kiểm tra phiên bản ILOM", ["", []]],
            [4, "Kiểm tra phiên bản Image", ["", []]],
            [5, "Kiểm tra cấu hình RAID và dung lượng phân vùng OS", ["", []]],
            [6, "Kiểm tra cấu hình Bonding Network", ["", []]],
            [7, "Kiểm tra CPU Utilization", ["", []]],
            # [8, "Kiểm tra CPU Load Average", ["", []]],
            [8, "Kiểm tra Memory", ["", []]],
            [9, "Kiểm tra IO Busy", ["", []]],
        ]
    elif system_info["type"] == "vm" \
            and system_info["platform"] == "solaris":
        checklist = [
            # ["STT", "Hạng mục kiểm tra", "Score"],
            [1, "Kiểm tra lỗi", ["", []]],
            [2, "Kiểm tra phiên bản Image", ["", []]],
            [3, "Kiểm tra dung lượng phân vùng OS", ["", []]],
            [4, "Kiểm tra CPU Utilization", ["", []]],
            [5, "Kiểm tra Memory", ["", []]],
            [6, "Kiểm tra IO Busy", ["", []]],
        ]

    keys = list(asserted)

    for i in range(0, len(checklist)):
        asserted_score = asserted[keys[i+1]][0]
        comment = asserted[keys[i+1]][1]
        try:
            if asserted_score == "":
                score = asserted_score
            else:
                score = ASSERTION[asserted_score]
        except Exception:
            score = asserted_score
        checklist[i][2][0] = score
        core.logger.info(checklist[i][1] + ":" + str(score))
        checklist[i][2][1] = comment

    return checklist


# This function table with last column cells may or may not contain
# list of string
def drw_table(doc: Document, checklist: list, row: int, col: int, info: bool = False):
    if checklist == []:
        return -1
    tab = doc.add_table(row, col)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    # tab.style = "Table Grid"

    # ADD TITLE CELLS AND COLOR THEM
    cols = tab.rows[0].cells
    for r in range(len(checklist[0])):
        cell = cols[r]
        cell.text = checklist[0][r]
        cell.paragraphs[0].style = "Table Heading"
        shading_elm = parse_xml(
            r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), TABLE_RED)
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # ADD CONTENT TO TABLE
    if not info:
        for i in range(1, len(checklist)):
            rows = tab.rows[i]
            cells = rows.cells
            for j in range(0, len(checklist[i])):
                cells[j].text = str(checklist[i][j])
                cells[j].paragraphs[0].style = "Table Paragraph"
    else:
        for i in range(1, len(checklist)):
            rows = tab.rows[i]
            cells = rows.cells
            for j in range(0, len(checklist[i])):
                if j == (len(checklist[i]) - 1):
                    cells[j].text = str(checklist[i][j][0])
                    cells[j].paragraphs[0].style = "Table Paragraph"
                    continue
                cells[j].text = str(checklist[i][j])
                cells[j].paragraphs[0].style = "Table Paragraph"
    return tab


# def drw_image_to_doc(doc, node, images_root, images_name):
def drw_info(
        doc: Document, node: str, checklist: list,
        images_root: Path, images_name: list = []
) -> None:
    prev = doc.add_paragraph("Thông tin chi tiết", style="baocao4")
    for i in range(0, len(checklist)):
        # doc.add_paragraph(checklist[i][1], style="baocao5")
        par = doc.add_paragraph(checklist[i][1])
        # list_number(doc, par, prev=prev, level=3)
        list_number(doc, par, prev=prev)
        try:
            if isinstance(images_name[i], list):
                for image in images_name[i]:
                    path = images_root  / node / image
                    doc.add_picture(str(path), width=Inches(6.27))
            else:
                path = images_root / node / images_name[i]
                doc.add_picture(
                    str(path),
                    width=Inches(6.27),
                )
        except Exception:
            pass
        for line in checklist[i][2][1]:
            doc.add_paragraph(line)
    doc.add_paragraph("")
    doc.add_page_break()


def define_doc(sample: Path):
    try:
        return docx.Document(sample)
    except Exception:
        print(f"{str(sample)} docx not found!")
        return None


def drw_menu(doc, nodes):
    # doc.add_paragraph("ORACLE EXADATA X8M-2", style="baocao1")
    # doc.add_paragraph("Kiểm tra nhiệt độ môi trường",
    #                   style=styles["baocao2"])
    doc.add_heading("Kiểm tra nhiệt độ môi trường", level=2)
    doc.add_paragraph("Mục lục", style="Heading")
    for node in nodes:
        doc.add_paragraph("Kiểm tra nhiệt độ môi trường",
                          style="baocao2")
        doc.add_paragraph("").paragraph_format.tab_stops.add_tab_stop(
            Inches(1.5), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS
        )
    doc.add_page_break()


system_info = {"system_type": "",
               "platform": "",
               "type": ""}


def drw_doc_appendix(
        doc: Document, checklist_list: list, nodes: list,
        images_root: Path, images_name: list
) -> None:
    for node in nodes:
        print("NODE:" + node["node_name"])
        print("RUNNING:GETTING SAVED IMAGES")
        image_json = images_root / (node["node_name"] + "/images.json")
        images_name = tools.read_json(image_json)
        print("RUNNING:DRAWING OVERVIEW TABLE")
        doc.add_paragraph(
            "Máy chủ " + node["node_name"], style="baocao3")
        # list_number(doc, par, prev=prev)
        # doc.add_heading(
        #     "Máy chủ " + node["node_name"], level=2)
        doc.add_paragraph("Thông tin tổng quát", style="baocao4")
        overview = [
            ["Hostname", "Product Name", "Serial Number", "IP Address"],
            [node["node_name"], "", "", ""],
        ]
        drw_table(doc, overview, 2, 4)
        doc.add_paragraph("")
        doc.add_paragraph("Đánh giá", style="baocao4")
        print("RUNNING:DRAWING SUMMARY TABLE")

        check_table = [
            ["STT", "Hạng mục kiểm tra", "Score"],
        ]
        check_table.extend(checklist_list[node["node_name"]])
        drw_table(doc, check_table, len(check_table), 3, True)
        doc.add_paragraph("")

        print("RUNNING:DRAWING DETAILS")
        drw_info(doc, node["node_name"],
                 checklist_list[node["node_name"]], images_root, images_name)
        # doc.add_page_break()

        print("DONE")
        print()


def drw_doc(doc, checklist_list, nodes):
    doc.add_paragraph("Máy chủ ?(S/N): ?", style="baocao2")

    print("RUNNING:DRAWING OVERVIEW TABLE")
    doc.add_paragraph("Thông tin chung", style="baocao4")
    overview = [
        ["ITEM", "VALUE", "VALUE(PREVIOUS REPORT)"],
    ]
    for node in nodes:
        print("NODE:" + node["node_name"])
        overview.extend([
            ["", "Máy chủ " + node["node_name"], ""],
            ["Hostname", node["node_name"], ""],
            ["Serial Number", "", ""],
            ["Image Version", "", ""],
            ["IP Adress", "", ""]
        ])
    # print(json.dumps(overview, indent=2))
    drw_table(doc, overview, len(overview), 3, False)
    doc.add_paragraph("")

    print("RUNNING:DRAWING DETAIL TABLES")
    doc.add_paragraph("Thông tin chi tiết", style="baocao4")
    for node in nodes:
        doc.add_paragraph(
            "Máy chủ " + node["node_name"], style="baocao3")
        detail = [
            ["STT", "Hạng Mục kiểm tra", "Điểm đánh giá",
                "Điểm đánh giá\n (Trong lần kiểm tra trước đây)"],
        ]
        for check in checklist_list[node["node_name"]]:
            detail.append([str(check[0]), str(check[1]), str(check[2][0]), ""])
        # print(json.dumps(detail))

        drw_table(doc, detail, len(detail), 4, False)
        doc.add_paragraph("")
    print("RUNNING:DRAWING RECOMMNEDATION TABLE")
    doc.add_paragraph("Khuyến cáo", style="baocao3")
    recommend = [
        ["No", "Khuyến cáo Rủi Ro", "Mức độ", "Note"],
        ["", "", "", ""]
    ]
    drw_table(doc, recommend, len(recommend), 4, False)
    doc.add_paragraph("")
    print("RUNNING:DRAWING REFERENCE")
    doc.add_paragraph(
        "Thông tin kiểm tra chi tiết cho hệ thống ?", style="baocao3")
    doc.add_paragraph(
        "Vui lòng kiểm tra tài liệu Appendix được gửi kèm báo cáo này.")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Ý KIẾN CÁC BÊN", style="baocao1")
    ref = [
        ["KHÁCH HÀNG", "MPS"],
        [" ", " "],
        ["Tên: ", "Tên: "],
        ["Chữ ký: ", "Chữ ký: "],
        ["Ngày: ", "Ngày: "],
    ]
    drw_table(doc, ref, len(ref), 2, False)
    print("DONE")
    print()


def compile(
        doc: Document, appendix_doc: Document, input_file: Path, out_dir: Path,
        images_root: Path, force: bool
) -> Document:
    input_file_data = tools.read_json(input_file)
    # Wrap in a list if input_file_data contains only 1 dictionary
    if not isinstance(input_file_data, list):
        input_file_data = [input_file_data]

    summary_content = []
    for time in range(len(input_file_data)):
        system_info["system_type"] = input_file_data[time]["system_type"]
        system_info["platform"] = input_file_data[time]["platform"]
        system_info["type"] = input_file_data[time]["type"]
        nodes = input_file_data[time]["nodes"]
        checklist_list = {}

        images_name = []
        if nodes == -1:
            raise RuntimeError(f"Nodes data not found in str{input_file}")

        asserted_list = []
        appendix_doc.add_page_break()
    # drw_menu(doc, nodes)

        # input_root = os.path.split(input_file)[0]
        input_root = input_file.parent
        print("RUNNING:ASSERTING DATA")
        for node in nodes:
            print("NODE:" + node["node_name"])
            asserted = assert_data(node)

            print("RUNNING:SAVING ASSERTED DATA")
            asserted_file = input_root / \
                node["node_name"] / "asserted.json"

            asserted_list += [asserted_file]
            tools.save_json(asserted_file, asserted, append=False)
            print("RUNNING:CREATING CHECKLIST")
            checklist = get_score(asserted)
            checklist_list[node["node_name"]] = checklist
            print("DONE")
            print()
        core.logger.debug(asserted_list)
        print("RUNNING:SAVING ASSERTED SUMMARY FILE")
        print()

        file_name = Path(tools.rm_ext(str(input_file), "json") +
                         f"_asserted-{time}.json")

        tools.save_json(
            file_name,
            system_info,
            False
        )

        tools.join_json(file_name, asserted_list)
        summary_content.append(tools.read_json(file_name))

        print("RUNNING:DRAWING REPORTS")
        print("RUNNING:DRAWING APPENDIX REPORT")
        if system_info["type"] == "baremetal" \
                and system_info["platform"] == "solaris":
            appendix_doc.add_paragraph(
                "Máy chủ SPARC", style="baocao1")
        elif system_info["type"] == "vm":
            appendix_doc.add_paragraph(
                "Máy chủ ảo hóa", style="baocao1")

        drw_doc_appendix(appendix_doc, checklist_list, nodes,
                         images_root, images_name)
        print("RUNNING:DRAWING REPORT")
        try:
            drw_doc(doc, checklist_list, nodes)
        except Exception:
            pass

    main_summary = Path(tools.rm_ext(
        str(input_file), "json") + "_asserted.json")
    tools.save_json(main_summary, summary_content, append=False)

    return appendix_doc


def print_style(doc):
    styles = doc.styles
    paragraphs_style = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
    # for style in styles:
    for style in paragraphs_style:
        print(style.name)


def run(input_file: Path, output_file: Path, sample: Path,
        appendix_sample: Path, images_dir: Path, force: bool = False) -> list:
    doc = define_doc(sample)
    appendix_doc = define_doc(appendix_sample)

    # out_dir = os.path.split(output_file)[0]
    out_dir = output_file.parent
    print_style(appendix_doc)
    try:
        appendix_doc = compile(doc, appendix_doc, input_file,
                               out_dir, images_dir, force)
        if appendix_doc == -1:
            return -1
    except Exception as err:
        print(err)
        print("Exiting")
        sys.exit()

    output_base_name = tools.rm_ext(output_file.name, "json")

    doc_name = ""
    if doc is not None:
        doc_name = str(out_dir / (output_base_name + ".docx"))
        doc.save(doc_name)

    appendix_doc_name = str(
        out_dir / ("appendix-" + output_base_name + ".docx"))
    appendix_doc.save(appendix_doc_name)

    return [doc_name, appendix_doc_name]


def main():
    run()


if __name__ == "__main__":
    main()
