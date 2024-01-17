#!/usr/bin/env python3


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import io
from PIL import Image
from PIL import ImageColor
from PIL import ImageDraw
from PIL import ImageFont
from docx.oxml import OxmlElement

from rekdoc.const import *
from rekdoc import tools

import uuid
# def drw_text_image(text, file):
#     with Image(width=1000, height=1000, background=Color("black")) as img:
#         img.format = "png"
#         with Drawing() as context:
#             tmp = text.getvalue()
#             metrics = context.get_font_metrics(img, tmp, True)
#             x = 10
#             y = 14
#             w = int(metrics.text_width * 1.2)
#             h = int(metrics.text_height * 1.3)
#             img.resize(w, h)
#             # context.font_family = "monospace"
#             context.font_size = y
#             context.fill_color = Color("white")
#             context.gravity = "north_west"
#             context.text_antialias = True
#             context.text(x, y, text.getvalue())
#             context(img)
#         img.save(filename="PNG24:" + file)


def drw_text_image(text, file):
    size = 14
    with Image.new("RGB", (1000, 1000)) as img:
        d1 = ImageDraw.Draw(img)
        left, top, right, bottom = d1.textbbox((10, 10), text, font_size=size)
        print(right)
        print(bottom)
        w = int(right * 1.1) + 10
        h = int(bottom * 1.1) + 10
        print(w)
        print(h)
        # img_resize = img.crop((0, 0, w, h))
        img_resize = img.resize((w, h), resample=Image.NEAREST)
        d2 = ImageDraw.Draw(img_resize)
        d2.text((10, 10), text, font_size=size)
        img_resize.save(file + ".png", format="PNG")


def set_font_size(run, font_size):
    font = run.font
    font.size = Pt(font_size)

def set_numbering_style(paragraph, level):
    numbering = paragraph._element.xpath('.//w:numPr')
    if numbering:
        num_id = numbering[0].xpath('@w:numId')[0]
        abstract_id = numbering[0].xpath('@w:ilvl')[0]
        numbering[0].remove(num_id)
        numbering[0].remove(abstract_id)

    paragraph._element.append(OxmlElement('w:numPr'))
    paragraph._element.xpath('.//w:numPr')[0].append(OxmlElement('w:ilvl', val=str(level)))
    paragraph._element.xpath('.//w:numPr')[0].append(OxmlElement('w:numId', val="1"))

def add_numbered_list(document, levels):
    for level, text in levels:
        paragraph = document.add_paragraph(text, style='BodyText')
        set_numbering_style(paragraph, level)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_custom_heading(document, text, level):
    heading_style = 'Heading' + str(level)
    paragraph = document.add_heading(text, level=level)
    paragraph.style = document.styles[heading_style]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Change the font size for the heading
    set_font_size(paragraph.runs[0], 14)

def add_subheading(document, heading_level, subheading_text):
    # Add a paragraph as a subheading under the specified heading level
    paragraph = document.add_paragraph()
    run = paragraph.add_run(subheading_text)
    set_numbering_style(paragraph, heading_level + 1)  # Increment the level for subheading
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

if __name__ == "__main__":
    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Document with Styled Headings and Numbered List', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Define the levels and texts for each level
    levels = [
        (1, 'Level 1'),
        (2, 'Level 2'),
        (3, 'Level 3'),
        (4, 'Level 4'),
        (5, 'Level 5')
    ]

    # Add custom headings
    for level, text in levels:
        add_custom_heading(doc, text, level)

        # Add a subheading for each main heading
        add_subheading(doc, level, f'Subheading under Level {level}')

    # Add the numbered list
    add_numbered_list(doc, levels)

    # Save the document
    doc.save('styled_document_with_subheadings.docx')
